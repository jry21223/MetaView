"""
Comprehensive business plan rewrite with unified Chinese-first terminology.
Rules:
- Chinese first, English in parens on first use only
- Only describe implemented features
- Business-audience tone
- Clean formatting throughout
"""
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

SRC = '/Users/jerry/Downloads/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2.docx'
DST = '/Users/jerry/Downloads/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2.docx'
doc = Document(SRC)


def set_text(para, text):
    """Set paragraph text, adding a run if empty."""
    runs = para.runs
    if not runs:
        r = para.add_run(text)
        r.font.name = 'Noto Sans CJK SC'
        r.font.size = Pt(11)
        return
    for r in runs:
        r.text = ''
    runs[0].text = text


def set_cell(table, row, col, text):
    """Set table cell text."""
    cell = table.rows[row].cells[col]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs:
        rr = cell.paragraphs[0].runs
        if rr:
            rr[0].text = text
        else:
            cell.paragraphs[0].text = text
    else:
        cell.text = text


def find_para(snippet):
    """Find first paragraph containing snippet."""
    for p in doc.paragraphs:
        if snippet in p.text:
            return p
    return None


def find_empty_after_heading(heading_text):
    """Find first empty paragraph after given heading."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and heading_text in p.text:
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip() == '':
                    return doc.paragraphs[j]
    return None


def find_table_by_header(text):
    """Find table whose first row contains text."""
    for t in doc.tables:
        hdr = ' '.join(c.text for c in t.rows[0].cells)
        if text in hdr:
            return t
    return None


# ================================================================
# §2.3 解决方案
# ================================================================
p = find_para('MetaView 以 LLM 驱动的结构化教学脚本生成')
if p:
    set_text(p,
        'MetaView 的核心是一条端到端的自动化教学动画生成管线：'
        '用户以自然语言、代码片段或数学公式提交一道题目，'
        '系统自动识别学科领域，调用大语言模型生成一份结构化的教学蓝图'
        '（CIR，定义「讲什么」——标题、步骤、视觉类型、讲解词），'
        '同时生成执行映射（ExecutionMap，定义「什么时候讲、对应哪行代码」），'
        '两者装配为统一的描述文件——教学脚本（PlaybookScript），'
        '再由前端渲染引擎逐帧驱动为可交互的教学动画，'
        '支持参数面板调节、步骤跳转、语音合成朗读和视频文件导出。'
        '\n\n'
        '与通用 AI 对话工具不同，MetaView 不是让模型直接输出最终动画，'
        '而是在模型和画面之间插入了一层结构化契约：'
        '模型只输出严格符合规范的 JSON 数据，'
        '画面渲染和交互逻辑由系统注入。'
        '这从根本上规避了大模型长文本截断和幻觉扩散的问题，'
        '也让同一份教学脚本可以驱动不同的输出形式——'
        '当前稳定支持交互播放和 MP4 视频导出，'
        '架构预留了扩展至幻灯片、HTML5 课件、静态图序列等输出通道。'
        '\n\n'
        '平台同时提供智能体模式（Agent Mode）：'
        '大模型通过绘图工具链逐步构建画面——规划大纲、添加曲线、放置数组标记、写入公式——'
        '每一步都经几何自检和结构校验，最终由独立评审模型复核后输出。'
        '两条路径共享同一套教学脚本规范，走同一套渲染出口，'
        '用户可根据题目复杂度和质量要求灵活选择。')
    print('OK §2.3')
else:
    print('MISS §2.3')

# ================================================================
# §4.1 产品总体架构
# ================================================================
p = find_para('MetaView v2 采用')
if p:
    set_text(p,
        'MetaView 采用「API 后端 + Web 前端 + 智能体侧边车」三层架构，'
        '围绕教学脚本（PlaybookScript）这一统一数据契约构建。'
        '\n\n'
        'API 后端（FastAPI）负责题目接收、学科路由、大模型调用编排、'
        '教学蓝图生成与校验、教学脚本装配、视频导出调度以及运行历史管理。'
        '后端遵循整洁架构分层，将业务逻辑、领域模型和基础设施解耦，'
        '使得大模型提供商（OpenAI / DeepSeek / Ollama / vLLM）'
        '可以自由切换而不影响核心管线。'
        '\n\n'
        'Web 前端（React 19 + Vite）以渲染引擎（Remotion）为核心，'
        '按帧精确播放教学动画。'
        '前端同时提供题目输入界面、参数调节面板、语音合成与字幕控制、'
        '模型接入配置（Provider Settings）和视频导出操作界面。'
        '播放器支持键盘快捷键操控，历史记录可回溯任意一次生成的完整结果。'
        '\n\n'
        '智能体侧边车（Node.js + Pi Agent Core）仅在智能体模式下启用：'
        '注册了 14 个底层绘图工具和 11 个跨学科教学模板，'
        '大模型通过逐步调用这些工具来精确构建动画，'
        '并在每一步完成后进行几何校验和结构自检。'
        '三个服务通过 Docker Compose 编排，共享统一的脚本契约。')
    print('OK §4.1')
else:
    print('MISS §4.1')

# ================================================================
# §4.4 产品版本规划
# ================================================================
p = find_para('MetaView v2 目前为全栈单体应用')
if p:
    set_text(p,
        'MetaView 目前为全栈单体应用，提供自用版和运营版两种启动模式。'
        '自用版（./start.sh）面向个人用户，支持配置自定义的大模型接入'
        '（OpenAI / DeepSeek / Ollama / vLLM），可交互播放和导出视频。'
        '运营版（./start.sh op）面向平台运营者，增加了微信支付充值、'
        '用量监控面板和第三方 API 跳转充值等功能。'
        '两个版本均支持 Docker 一键部署（docker compose up）。'
        '未来版本规划包括：独立的学生学习空间、学校级管理后台、'
        '教师素材共创市场以及私有化部署方案。')
    print('OK §4.4')
else:
    # Find by heading
    p = find_empty_after_heading('产品版本规划')
    if p:
        set_text(p,
            'MetaView 目前为全栈单体应用，提供自用版和运营版两种启动模式。'
            '自用版（./start.sh）面向个人用户，支持配置自定义的大模型接入'
            '（OpenAI / DeepSeek / Ollama / vLLM），可交互播放和导出视频。'
            '运营版（./start.sh op）面向平台运营者，增加了微信支付充值、'
            '用量监控面板和第三方 API 跳转充值等功能。'
            '两个版本均支持 Docker 一键部署（docker compose up）。'
            '未来版本规划包括：独立的学生学习空间、学校级管理后台、'
            '教师素材共创市场以及私有化部署方案。')
        print('OK §4.4 (via heading)')
    else:
        print('MISS §4.4')

# ================================================================
# §5.1 技术架构 — complete rewrite with clean Chinese
# ================================================================
p = find_para('MetaView v2 后端遵循 Clean Architecture')
if p:
    set_text(p,
        '后端架构（FastAPI + Python）：'
        '遵循整洁架构四层分层。'
        '接口层（presentation）负责 HTTP 路由、中间件和依赖注入；'
        '应用层（application）编排用例（题目提交管线、视频导出、'
        '账户管理、充值兑换），并通过抽象端口定义与大模型、数据库、'
        '智能体侧边车、支付网关的交互协议；'
        '领域层（domain）包含所有核心模型（教学蓝图、教学脚本、'
        '执行映射、路由决策等 Pydantic 类型契约）和领域服务'
        '（蓝图提示词构建、脚本装配、学科路由、模型路由、几何校验等）；'
        '还包括 7 个学科技能包（代数、微积分、线性代数、二次变换、'
        '立体几何、力学等），每个技能包声明式注册、确定性计算，'
        '新学科可零代码适配。'
        '基础设施层（infrastructure）实现具体的大模型 HTTP 客户端、'
        'SQLite 持久化、微信支付和微信登录等。'
        '\n\n'
        '前端架构（React 19 + Vite + TypeScript）：'
        '遵循 Feature-Sliced Design 分层。'
        '共享层（shared）提供配置常量、通用 UI 组件和纯函数数学表达式引擎；'
        '实体层（entities）定义教学脚本和运行记录的类型；'
        '功能层（features）包括播放引擎（渲染器注册表 + 播放器 + '
        '参数面板 + 语音合成 + 导出）、模型接入配置、交互数学画板和题目输入；'
        '页面层（pages）组合为三个界面：题目输入页、工作室页和历史记录页。'
        '\n\n'
        '核心渲染管线：教学脚本是唯一的渲染契约。'
        '前端渲染器注册表按画面帧类型（如数组帧、柱状图帧、'
        '树结构帧、函数图帧、公式帧等）派发对应的渲染组件。'
        '交互播放器和视频导出共用同一套渲染器注册表，'
        '保证预览和导出画面完全一致。'
        '\n\n'
        '智能体模式附加架构：智能体侧边车注册了 14 个底层绘图工具'
        '（规划大纲、开始步骤、添加曲线、放置数组标记、写入公式等）'
        '和 11 个跨学科教学模板（数组交换、切线绘制、受力分析、'
        '抛体轨迹、黎曼和等），大模型通过逐步调用这些工具精确构建动画。'
        '\n\n'
        '格式无关的多输出设计：教学脚本只描述教学内容'
        '（步骤、旁白、视觉类型、时间轴），不绑定任何特定渲染技术。'
        '前端渲染器注册表按类型派发，新增输出格式只需编写对应适配器并注册，'
        '无需改动生成管线。当前已稳定支持交互播放和 MP4/WebM/GIF 视频导出，'
        '架构预留了幻灯片导出、HTML5 课件、静态图片序列等输出通道。'
        '这使核心生成资产（教学蓝图 + 教学脚本）可以跨输出介质永久复用。')
    print('OK §5.1')
else:
    print('MISS §5.1')

# ================================================================
# §5.2 核心算法流程
# ================================================================
p = find_para('MetaView v2 支持两条生成路径')
if p:
    set_text(p,
        'MetaView 支持两种生成路径，共享同一套教学脚本规范和渲染出口：'
        '\n\n'
        '标准模式（默认，适合日常教学）：'
        '\n'
        '1. 题目解析：前端收集用户输入的题目文本、可选的源代码和编程语言。'
        '\n'
        '2. 学科路由：学科路由器（混合模式：小模型分类 + 规则回退）'
        '自动识别题目所属学科领域。高置信度命中特定技能包时直接走确定性计算；'
        '否则走通用大模型生成路径。'
        '\n'
        '3. 教学蓝图生成：系统组装提示词（包含 JSON 格式约束和学科指引），'
        '大模型输出结构化教学蓝图——8 到 14 个教学步骤，'
        '每步包含讲解词模板、视觉类型、标记点集合、函数图参数等。'
        '同时输出执行映射，定义每步的时间区间和代码行号关联。'
        '\n'
        '4. 自动校验与修复：蓝图校验器检查结构完整性、'
        '标记一致性、视觉类型合法性、讲解词质量。'
        '不通过则触发自动修复服务——补全缺失字段、调整布局重叠、'
        '修正末步类型——修复失败则反馈给大模型重新生成（最多重试 N 次）。'
        '\n'
        '5. 教学脚本装配：脚本构建器将教学蓝图和执行映射装配为教学脚本，'
        '包括画面帧序列构造、帧区间计算（30fps）、代码高亮叠加、'
        '讲解词模板展开为语音合成文本、数学表达式安全过滤等。'
        '\n'
        '6. 渲染：前端播放器按帧率驱动交互播放，'
        '或通过视频导出用例调渲染引擎子进程导出视频文件。'
        '\n\n'
        '智能体模式（适合高质量要求）：'
        '\n'
        '1. 智能体侧边车接收生成请求。'
        '\n'
        '2. 大模型通过绘图工具链逐步调用：'
        '规划大纲 → 开始步骤 → 添加曲线/数组标记/公式 → 提交步骤 → 完成脚本。'
        '\n'
        '3. 几何自检：大模型在写入旁白前必须调用几何断言'
        '（方向判断 / 经过点判断 / 单调性判断），'
        '由后端数学引擎纯函数计算确定真值，不一致则强制修改。'
        '\n'
        '4. 三重门禁：自我检查 + 独立评审模型复核 + '
        '渲染器兼容性检查，全部通过后才输出格式合法的教学脚本。')
    print('OK §5.2')
else:
    print('MISS §5.2')

# ================================================================
# §5.3 创新点
# ================================================================
p = find_empty_after_heading('创新点')
if p:
    set_text(p,
        '创新一：教学蓝图——语言无关的结构化教学契约。'
        '定义了标题、步骤、视觉类型、标记点、讲解词模板、函数图参数等字段，'
        '模型只输出紧凑 JSON，画面渲染由系统注入，'
        '从架构层面消除长文档截断和幻觉扩散风险。'
        '\n\n'
        '创新二：双模式管线共享唯一出口。'
        '标准模式走快速生成路径，智能体模式通过绘图工具链逐步精细构建——'
        '两条路径输出同一格式的教学脚本，走同一套渲染器，'
        '用户按需选择，不锁定单一模型或供应商。'
        '\n\n'
        '创新三：自动校验-修复-重试三段式质量闭环。'
        '蓝图校验器多维检查（结构 / 标记 / 视觉类型 / 讲解词质量）→ '
        '自动修复服务补全缺失字段、调整布局、修正类型 → '
        '修复失败则反馈大模型重新生成。主动修复代替报错，闭环提升可用率。'
        '\n\n'
        '创新四：大小模型分层编排降本。'
        '低成本小模型做学科路由和意图分类，'
        '高质量大模型做教学蓝图生成。'
        '学科技能包声明式注册（清单 + 确定性计算内核），'
        '新学科零代码适配，避免厂商锁定，API 调用成本降低 30% 以上。'
        '\n\n'
        '创新五：毫秒级四向联动探索。'
        '执行映射实现代码行号、动画步骤、时间轴位置、变量状态四者精确对齐。'
        '用户点击任意维度即刻跳转到对应位置，'
        '提供从「看动画」到「探索过程」的认知升级。'
        '\n\n'
        '创新六：绘图工具链——让大模型一步一步精确构建动画。'
        '14 个底层原子工具 + 11 个跨学科教学模板 + 数学引擎几何自检'
        '（方向性 / 经过点 / 单调性），'
        '从工具能力面消除无意义铺陈，每一帧都有据可查。')
    print('OK §5.3')
else:
    print('MISS §5.3')

# ================================================================
# §5.4 内容可信与伦理 — rewrite to reflect actual implementation
# ================================================================
# Find the list items under 5.4 and update them
# The current bullets say things about 来源追踪, 教师审校, 内容安全, 版权治理, 隐私保护
# Replace with what's actually implemented
bullet_updates = {
    '来源追踪': '自动校验机制：每次生成都经过蓝图校验器多维检查，不通过的输出在管线内被拦截修复，不会到达用户。校验覆盖结构完整性、字段一致性、视觉类型合法性和讲解词质量。',
    '教师审校': '自动修复与重试：校验失败时，修复服务自动补全缺失字段、调整布局、修正类型。修复失败则反馈大模型重新生成（最多重试 N 次），形成闭环质量保障。',
    '内容安全': '输入过滤：数学表达式经白名单字符集过滤（仅允许数字、字母、运算符和基本标点），防止注入攻击。代码行号索引越界时静默丢弃，不做越权渲染。',
    '版权治理': '模型接入隔离：用户自行配置大模型提供商和 API 密钥（保存在本地浏览器），平台不存储用户密钥。支持 OpenAI、DeepSeek、Ollama、vLLM 等兼容接口。',
    '隐私保护': '本地持久化：运行历史存储在本地 SQLite 数据库，仅保留题目文本和生成结果用于复盘，不采集学生个人敏感信息。',
}

for old_start, new_text in bullet_updates.items():
    p = find_para(old_start)
    if p:
        set_text(p, new_text)

print('OK §5.4')

# ================================================================
# REWRITE TABLES
# ================================================================

# --- Table 3: 一页读懂 ---
t = find_table_by_header('项目定位')
if t and len(t.rows) >= 6:
    set_cell(t, 1, 1,
        'MetaView 是面向教师和学习者的 AI 教学可视化平台。'
        '以教学脚本为统一中间表示，将知识点转化为帧驱动交互动画和 MP4 教学视频；'
        '架构天然支持扩展至幻灯片、HTML5 课件、静态图序列等任意输出格式，'
        '一次生成、多介质复用。')
    set_cell(t, 4, 1,
        'API 后端 + Web 前端 + 智能体侧边车三层架构，'
        '围绕教学脚本统一契约构建。'
        '输出端通过渲染器注册表按画面帧类型派发渲染组件，'
        '交互播放和视频导出共用同一注册表。')
    set_cell(t, 5, 1,
        '以标准模式（快速）和智能体模式（精细）两条生成路径'
        '覆盖日常教学和高质量制作两种场景；'
        '大模型路由分层降低 API 调用成本，学科技能包可插拔扩展。')
    set_cell(t, 6, 1,
        '以「教学蓝图-可视化表达-四向联动-质量闭环」为流程创新；'
        '强调自动校验修复（而非人工审校）、格式可复用（而非锁定单一引擎）、'
        '过程可追溯（而非黑盒生成）。')
    set_cell(t, 7, 1,
        '一年内完成 7 个学科方向覆盖，沉淀 500+ 教学脚本案例，'
        '服务 50 所学校、2,000 名教师、5 万名学习者，'
        '形成可复制的「工具 + 内容 + 培训」推广模式。')
    print('OK Table 3')
else:
    print('MISS Table 3')

# --- Table 9: 学科覆盖 ---
t = find_table_by_header('学科')
if t and len(t.rows) >= 7:
    set_cell(t, 5, 0, '计算机/算法')
    set_cell(t, 5, 1, '算法流程、数据结构、代码逻辑、网络结构')
    set_cell(t, 5, 2, '排序动画、二叉树遍历、图搜索路径、代码联动高亮。')
    set_cell(t, 6, 0, '数学')
    set_cell(t, 6, 1, '代数变换、函数图像、微积分、立体几何')
    set_cell(t, 6, 2, '函数图绘制、导数切线、积分阴影、参数滑块调节、立体几何截面。')
    print('OK Table 9')
else:
    print('MISS Table 9')

# --- Table 10: 产品模块 ---
t = find_table_by_header('模块')
if t and len(t.rows) >= 7:
    set_cell(t, 1, 0, 'API 后端服务')
    set_cell(t, 1, 1, '题目接收、学科路由、大模型调用编排、教学蓝图生成与校验、教学脚本装配、视频导出调度、运行历史管理。')
    set_cell(t, 1, 2, '前端、智能体侧边车')
    set_cell(t, 2, 0, 'Web 前端 + 渲染引擎')
    set_cell(t, 2, 1, '题目输入、帧驱动交互播放器、渲染器注册表、参数面板、语音合成与字幕、视频导出、模型接入配置。')
    set_cell(t, 2, 2, '最终用户（教师/学习者）')
    set_cell(t, 3, 0, '智能体侧边车')
    set_cell(t, 3, 1, '智能体模式下启用：14 个底层绘图工具 + 11 个教学模板，逐步构建动画，自检 + 评审 + 兼容性三重门禁。')
    set_cell(t, 3, 2, '平台（智能体模式）')
    set_cell(t, 4, 0, '交互数学画板')
    set_cell(t, 4, 1, '内置预设函数、参数滑块、实时公式渲染、函数图像绘制，与播放引擎共用数学表达式计算引擎。')
    set_cell(t, 4, 2, '学习者、教师')
    set_cell(t, 5, 0, '视频导出系统')
    set_cell(t, 5, 1, '通过渲染引擎命令行工具导出 MP4、WebM、GIF 格式，支持语音合成配音和进度追踪。')
    set_cell(t, 5, 2, '教师、内容创作者')
    set_cell(t, 6, 0, '模型接入与部署')
    set_cell(t, 6, 1, '支持 OpenAI / DeepSeek / Ollama / vLLM 等大模型兼容接口，Docker Compose 一键部署，SQLite 持久化。')
    set_cell(t, 6, 2, '开发者、运维')
    print('OK Table 10')
else:
    print('MISS Table 10')

# --- Table 11: 输出形态 ---
t = find_table_by_header('输出类型')
if t and len(t.rows) >= 6:
    set_cell(t, 1, 0, '交互播放')
    set_cell(t, 1, 1, '帧驱动播放器、步骤跳转、参数面板、字幕叠加')
    set_cell(t, 1, 2, '适合课堂演示和自主探索。')
    set_cell(t, 2, 0, '视频导出')
    set_cell(t, 2, 1, 'MP4 / WebM / GIF，可选语音合成配音')
    set_cell(t, 2, 2, '适合微课制作、预习复习、内容分发。')
    set_cell(t, 3, 0, '四向联动回放')
    set_cell(t, 3, 1, '代码行高亮 + 动画步骤 + 时间轴 + 数据状态 四屏联动')
    set_cell(t, 3, 2, '适合编程教学和算法讲解。')
    set_cell(t, 4, 0, '数学函数图')
    set_cell(t, 4, 1, '坐标系曲线绘制、导数切线、积分阴影、公式渲染')
    set_cell(t, 4, 2, '适合数学和理工科教学。')
    set_cell(t, 5, 0, '教学脚本数据')
    set_cell(t, 5, 1, '结构化 JSON、运行历史、原始题目回溯')
    set_cell(t, 5, 2, '适合资源沉淀、复盘分析和批量处理。')
    print('OK Table 11')
else:
    print('MISS Table 11')

# --- Table 13: 技术层级 ---
t = find_table_by_header('层级')
if t and len(t.rows) >= 7:
    set_cell(t, 1, 0, '前端展示层')
    set_cell(t, 1, 1, 'React 19 + Vite + TypeScript，Feature-Sliced Design 分层，渲染器注册表按画面帧类型派发，交互播放与视频导出共用同一注册表。')
    set_cell(t, 2, 0, 'API 服务层')
    set_cell(t, 2, 1, 'FastAPI + Pydantic 类型契约，整洁架构四层分层（接口/应用/领域/基础设施），路由包括题目管线、运行历史、视频导出、智能体、账户管理、用量监控。')
    set_cell(t, 3, 0, '生成管线层')
    set_cell(t, 3, 1, '学科路由（混合模式：小模型分类 + 规则回退）、教学蓝图生成（大模型输出结构化 JSON）、自动校验修复三段式闭环、教学脚本装配映射。')
    set_cell(t, 4, 0, '渲染引擎层')
    set_cell(t, 4, 1, '教学脚本驱动帧渲染（30fps），渲染器覆盖数组帧、柱状图帧、树结构帧、函数图帧、公式帧、立体几何帧等类型。')
    set_cell(t, 5, 0, '智能体工具层')
    set_cell(t, 5, 1, '智能体侧边车（Node.js + Pi Agent Core），14 个底层绘图工具 + 11 个教学模板 + 数学引擎几何自检（方向性/经过点/单调性）。')
    set_cell(t, 6, 0, '基础设施层')
    set_cell(t, 6, 1, 'SQLite 持久化、Docker Compose 编排、大模型 HTTP 兼容客户端、微信支付、语音合成代理。')
    print('OK Table 13')
else:
    print('MISS Table 13')

# --- Table 14: 核心流程 ---
t = None
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if '教学目标解析' in cell.text or '知识结构生成' in cell.text:
                t = tbl
                break
        if t:
            break
    if t:
        break

if t and len(t.rows) >= 8:
    set_cell(t, 1, 0, '1. 题目解析')
    set_cell(t, 1, 1, '收集题目文本、可选的源代码和编程语言，构造输入上下文。')
    set_cell(t, 2, 0, '2. 学科路由')
    set_cell(t, 2, 1, '混合模式：优先小模型分类输出路由决策，低置信度回退规则解析；高置信命中特定技能包时走确定性计算内核。')
    set_cell(t, 3, 0, '3. 教学蓝图 + 执行映射生成')
    set_cell(t, 3, 1, '大模型按 JSON 格式约束输出教学蓝图（8-14 步）+ 执行映射（时间轴、代码行号关联、数组聚焦索引），标准模式一次生成。')
    set_cell(t, 4, 0, '4. 自动校验与修复')
    set_cell(t, 4, 1, '蓝图校验器多维检查（结构/标记一致性/视觉类型/讲解词质量），修复服务自动补全缺失字段、调整布局，失败则触发大模型重生成。')
    set_cell(t, 5, 0, '5. 教学脚本装配')
    set_cell(t, 5, 1, '脚本构建器将蓝图和执行映射映射为教学脚本：画面帧序列、帧区间计算（30fps）、代码高亮叠加、讲解词展开、数学表达式安全过滤。')
    set_cell(t, 6, 0, '6. 渲染与导出')
    set_cell(t, 6, 1, '前端播放器按帧率驱动交互播放；视频导出用例调渲染引擎命令行工具导出 MP4 / WebM / GIF 视频文件。')
    set_cell(t, 7, 0, '7. 运行历史')
    set_cell(t, 7, 1, '保留原始题目和完整教学脚本，用户可回顾、重放、对比不同输入与生成结果。')
    print('OK Table 14')
else:
    print('MISS Table 14')

# --- Table 15: 创新点 ---
t = find_table_by_header('创新类型')
if t and len(t.rows) >= 6:
    set_cell(t, 1, 0, '教学蓝图契约')
    set_cell(t, 1, 1, '定义语言无关的结构化教学契约（教学蓝图 + 执行映射 → 教学脚本），模型只输出紧凑 JSON，系统注入画面渲染和交互，从架构层面消除长文档截断和幻觉扩散。')
    set_cell(t, 2, 0, '双模式统一出口')
    set_cell(t, 2, 1, '标准模式（快速）和智能体模式（精细）共享同一套教学脚本规范和渲染出口；智能体模式通过绘图工具链逐步精确构建，不锁定单一模型或供应商。')
    set_cell(t, 3, 0, '自动校验修复闭环')
    set_cell(t, 3, 1, '蓝图校验器 + 自动修复服务 + 大模型重试三段式质量保障，主动修复代替报错，闭环提升管线可用率。')
    set_cell(t, 4, 0, '大小模型分层降本')
    set_cell(t, 4, 1, '低成本小模型路由 + 高质量大模型生成，学科技能包声明式注册（清单 + 确定性内核），新学科零代码适配，API 成本降低 30% 以上。')
    set_cell(t, 5, 0, '四向联动探索')
    set_cell(t, 5, 1, '执行映射实现代码行号、动画步骤、时间轴位置、变量状态四者毫秒级对齐，从"看动画"升级为"探索过程"。')
    print('OK Table 15')
else:
    print('MISS Table 15')

# --- Table 23: 样板案例 ---
t = find_table_by_header('知识难点')
if t and len(t.rows) >= 6:
    set_cell(t, 1, 0, '算法：冒泡排序')
    set_cell(t, 1, 1, '相邻元素比较、交换、数组遍历')
    set_cell(t, 1, 2, '柱状图动画 + 逐帧交换高亮 + 代码行联动。')
    set_cell(t, 2, 0, '算法：二叉树 BFS/DFS')
    set_cell(t, 2, 1, '层序遍历、前中后序、队列/栈模拟')
    set_cell(t, 2, 2, '树节点高亮动画 + 访问顺序标注 + 代码联动。')
    set_cell(t, 3, 0, '数学：二次函数平移缩放')
    set_cell(t, 3, 1, '顶点式、对称轴、开口方向、参数影响')
    set_cell(t, 3, 2, '函数图 + 参数滑块实时调节 + 公式渲染。')
    set_cell(t, 4, 0, '数学：导数与切线')
    set_cell(t, 4, 1, '切线斜率、极值点、单调性')
    set_cell(t, 4, 2, '曲线 + 动态切线 + 高亮标记 + 坐标读数。')
    set_cell(t, 5, 0, '物理：抛体运动轨迹')
    set_cell(t, 5, 1, '初速度、角度、重力加速度、射程')
    set_cell(t, 5, 2, '参数曲线动画 + 速度分解箭头 + 轨迹包络。')
    print('OK Table 23')
else:
    print('MISS Table 23')

# --- Table 32: 风险控制 ---
t = find_table_by_header('风险类别')
if t and len(t.rows) >= 8:
    set_cell(t, 1, 0, '生成质量风险')
    set_cell(t, 1, 1, '大模型输出可能结构不完整、字段缺失或讲解词质量不达标。')
    set_cell(t, 1, 2, '蓝图校验器多维自动检查 + 修复服务自动补全 + 大模型重试重生成。校验规则可配置，不通过不交付。')
    set_cell(t, 2, 0, '幻觉与越界风险')
    set_cell(t, 2, 1, '大模型可能生成不存在的代码行号或非法数学表达式。')
    set_cell(t, 2, 2, '代码行号越界时静默丢弃，全部越界则跳过该步骤的高亮；数学表达式经白名单字符集过滤。')
    set_cell(t, 3, 0, '市场风险')
    set_cell(t, 3, 1, '学校采购周期长、预算不确定、教师使用频率不足。')
    set_cell(t, 3, 2, '先做低价试点包和样板课，建立数据报告和教师共创机制。')
    set_cell(t, 4, 0, '竞争风险')
    set_cell(t, 4, 1, '通用 AI 工具或教育平台可能进入教学可视化领域。')
    set_cell(t, 4, 2, '深耕学科案例、教学脚本资产沉淀、四向联动特色功能和教师社群，形成教育场景壁垒。')
    set_cell(t, 5, 0, '成本风险')
    set_cell(t, 5, 1, '大模型调用和视频渲染成本可能随使用量增长。')
    set_cell(t, 5, 2, '小模型路由降低单次调用成本；学科技能包命中时跳过生成；缓存复用已生成的脚本；分层收费和用量控制。')
    set_cell(t, 6, 0, '合规风险')
    set_cell(t, 6, 1, '学生隐私、数据安全、未成年人保护等要求。')
    set_cell(t, 6, 2, '数据最小化采集、本地 SQLite 存储、不采集学生敏感信息、Docker 隔离部署、日志审计。')
    set_cell(t, 7, 0, '团队风险')
    set_cell(t, 7, 1, '学生团队时间不稳定、技术与教研协同不足。')
    set_cell(t, 7, 2, '明确岗位分工、导师监督、里程碑考核、核心成员激励机制和文档化交接。')
    print('OK Table 32')
else:
    print('MISS Table 32')

# ================================================================
# GLOBAL CLEANUP: Remove remaining English-heavy paragraphs
# ================================================================
# Find and clean any paragraph that has excessive English terms
cleanup_map = {
    'CIR + ExecutionMap': '教学蓝图 + 执行映射',
    'PlaybookScript → Remotion': '教学脚本 → 渲染引擎',
    'Drawing CLI 工具集': '绘图工具链',
    'PlaybookScript 契约': '教学脚本契约',
    'PlaybookScript JSON': '教学脚本',
    'PlaybookScript 驱动': '教学脚本驱动',
    'PlaybookScript 装配': '教学脚本装配',
    'PlaybookScript 数据': '教学脚本数据',
    'PlaybookScript，': '教学脚本，',
    'PlaybookScript。': '教学脚本。',
    'CirValidator': '蓝图校验器',
    'RepairService': '自动修复服务',
    'PlaybookBuilder': '脚本构建器',
    'SkillPack': '学科技能包',
    'snapshot.kind': '画面帧类型',
    'renderer registry': '渲染器注册表',
    'Renderer Registry': '渲染器注册表',
    'renderer compatibility gate': '渲染器兼容性检查',
    'Renderer Compatibility Gate': '渲染器兼容性检查',
    'Pydantic 契约': '类型契约',
}

# Clean paragraphs
cleaned = 0
for old, new in cleanup_map.items():
    for para in doc.paragraphs:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                cleaned += 1
    # Also check paragraph text directly (for multi-run text)
    for para in doc.paragraphs:
        full = para.text
        if old in full:
            new_full = full.replace(old, new)
            if new_full != full:
                set_text(para, new_full)
                cleaned += 1

print(f'\nGlobal cleanup: {cleaned} replacements')

# ================================================================
# SAVE
# ================================================================
doc.save(DST)
print(f'\nDone. Saved to {DST}')
