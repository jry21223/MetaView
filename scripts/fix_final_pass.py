"""
Final comprehensive fix per user requirements:
- §2.2 expand
- §2.3 restructure (tech summary → pain points → future)
- §4.1 make layperson-friendly
- §4.2 natural Chinese
- §5.1/5.2/5.3 tone down from dev docs
- Tables: clean prohibited terms
"""
from docx import Document
from docx.shared import Pt, Cm

SRC = '/Users/jerry/Downloads/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2.docx'
DST = '/Users/jerry/Downloads/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2.docx'
doc = Document(SRC)


def set_text(para, text):
    runs = para.runs
    if not runs:
        r = para.add_run(text)
        r.font.name = 'Noto Sans CJK SC'
        r.font.size = Pt(11)
        return
    for r in runs:
        r.text = ''
    runs[0].text = text


def set_cell(table, row, col, text):
    cell = table.rows[row].cells[col]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs:
        rr = cell.paragraphs[0].runs
        if rr: rr[0].text = text
        else: cell.paragraphs[0].text = text
    else:
        cell.text = text


# ====================================================================
# §2.2 项目背景 — expand from 2 paragraphs
# ====================================================================
# Para 52: current "教育数字化和生成式人工智能..." → expand
p52 = doc.paragraphs[52]
set_text(p52,
    '教育数字化浪潮与国家政策推动：'
    '教育部等多部门持续推进教育数字化战略，'
    '强调以人工智能助力教育变革、扩大优质教育资源覆盖面。'
    '学校已普遍配备多媒体教室和在线教学平台，'
    '但在"用什么内容来填充这些数字化设施"这个关键环节上，'
    '仍然高度依赖教师个人制作——而教师最缺的恰恰是时间。'
    '\n\n'
    '教师的真实困境：'
    '备一节含图解、动画或交互演示的课，教师往往需要花费数小时甚至数天。'
    '算法、函数变换、物理过程、化学反应机理等抽象内容，'
    '用静态 PPT 和板书极难讲透。'
    '市面上出现了一些 AI 课件工具，但多数只是生成文字大纲或静态图片，'
    '缺乏将抽象概念"演绎出来"的动态能力，'
    '更无法做到让学生与内容双向互动。'
    '\n\n'
    '学习者的需求变化：'
    '今天的学生成长于短视频和交互应用时代，'
    '对静态文字和图片的耐心持续下降。'
    '他们需要的不是"看答案"，而是"看过程"——'
    '看到算法的每一步交换、看到函数图像如何随参数变化、'
    '看到物理受力分析的动态演变。'
    '更进一步，他们需要能在动画中点击、拖动、修改参数，'
    '从"被动观看"升级为"主动探索"。'
    '\n\n'
    '技术窗口已打开：'
    '大语言模型的成熟和前端渲染引擎的进步，'
    '使得"自然语言描述教学意图 → 自动生成结构化教学动画"'
    '这一技术路线从不可能变为可行。'
    '教育领域正是这一能力的最佳应用场景之一，'
    '但目前市场上尚无产品将大模型生成能力与帧精确渲染引擎'
    '以结构化契约的方式整合为完整教学管线。')

# Para 53: clear and remove (content merged into p52)
set_text(doc.paragraphs[53], '')
print('OK §2.2 expanded')

# ====================================================================
# §2.3 解决方案 — restructure as 3 parts
# ====================================================================
set_text(doc.paragraphs[55],
    '技术概要：'
    'MetaView 的核心是一条自动化教学动画管线。'
    '用户以自然语言、代码片段或数学公式输入一道题目，'
    '系统自动识别学科领域后，调用大语言模型生成一份结构化的教学蓝图'
    '（定义讲什么：标题、分步讲解、视觉类型、标记点），'
    '同时生成执行映射（定义何时讲、对应哪行代码），'
    '两者装配为统一的教学脚本，由前端渲染引擎逐帧驱动为可交互的动画。'
    '整个过程约 30 到 45 秒，无需人工干预。'
    '\n\n'
    '与通用 AI 工具的关键区别在于，MetaView 在模型输出和最终画面之间'
    '插入了一层结构化契约：模型只输出符合严格规范的 JSON 数据，'
    '画面渲染和交互逻辑完全由系统控制。'
    '这从根本上规避了大模型长文本截断和幻觉扩散问题——'
    '不通过自动校验的脚本不会进入渲染环节，'
    '结构缺失会自动修复，修复失败则让模型重试。'
    '\n\n'
    '未来展望：'
    '当前平台已稳定支持 7 个学科领域的交互播放和 MP4 视频导出。'
    '教学脚本是格式无关的描述文件，'
    '同一个脚本未来可驱动不同输出形式——'
    '幻灯片课件、HTML5 交互页面、静态图片序列等，无需改动生成管线。'
    '智能体模式（通过绘图工具链逐步精确构建）'
    '则为高质量教学视频制作提供了更精细的控制路径。')

print('OK §2.3 restructured')

# ====================================================================
# §4.1 产品总体架构 — make layperson-friendly
# ====================================================================
set_text(doc.paragraphs[76],
    'MetaView 是一个 Web 应用，用户打开浏览器即可使用，无需安装任何软件。'
    '界面分为三个主要区域：'
    '\n\n'
    '题目输入页——这就是"告诉系统你想讲什么"的地方。'
    '可以输入一段文字描述、粘贴代码、写数学公式，'
    '也可以上传参考图片。支持 Markdown 格式。'
    '\n\n'
    '工作室页——这是核心的工作区。'
    '左侧展示原始题目和 AI 对话面板（可直接与大模型交流调整内容），'
    '右侧是教学动画播放器。'
    '播放器不只是"播放"，而是一个交互探索环境：'
    '动画、代码高亮、时间轴、数据状态四者联动——'
    '点击任意维度，其他三个维度同步跳转到对应位置。'
    '同时提供参数调节面板（修改输入参数即时看到变化）、'
    '语音合成开关、字幕显示和播放速度控制。'
    '\n\n'
    '历史记录页——所有生成过的动画都保留在这里，'
    '包含原始题目和完整脚本数据，可以随时回放、对比和复盘。'
    '\n\n'
    '用户可以在设置中配置自己的大模型接入信息'
    '（支持 OpenAI、DeepSeek、本地 Ollama 或 vLLM 等兼容接口），'
    '密钥保存在本地浏览器中。'
    '视频导出功能可以将任意教学动画输出为 MP4、WebM 或 GIF 文件。'
    '平台也提供 Docker 一键部署方案，适合学校内网环境。')

print('OK §4.1 simplified')

# ====================================================================
# §4.2 教师端核心流程 — natural Chinese, no code names in parens
# ====================================================================
steps_42 = [
    (79,
     '1. 输入题目。'
     '用户在输入页面以自然语言、代码片段或数学公式描述题目。'
     '如果涉及编程教学，可以粘贴源代码，系统会自动关联代码行和动画步骤。'),
    (80,
     '2. 自动识别学科。'
     '系统根据题目内容自动判断所属学科领域'
     '（当前覆盖算法、数学、代码、物理、化学、生物、地理七个方向），'
     '并匹配合适的教学策略和步骤规划。'),
    (81,
     '3. 生成教学蓝图。'
     '大语言模型根据题目和学科，输出一份结构化的教学方案：'
     '包括教学步骤划分、每步的视觉表现形式、'
     '关键数据标记点、分步讲解词，以及每一步的时间安排。'),
    (82,
     '4. 自动校验与装配。'
     '系统对生成的教学方案进行多维检查——'
     '结构是否完整、标记是否一致、讲解是否达标。'
     '发现缺失或错误自动修复；修复不了则让模型重新生成。'
     '校验通过后装配为完整教学脚本。'),
    (84,
     '5. 交互播放与导出。'
     '前端渲染引擎按每秒 30 帧驱动教学动画。'
     '支持步骤跳转、参数面板调节、语音朗读和字幕显示。'
     '可以导出为 MP4、WebM 或 GIF 视频文件。'),
    (85,
     '6. 历史沉淀与复盘。'
     '每次生成自动保存原始题目和完整脚本。'
     '用户可以随时回看历史、重放动画、'
     '修改题目重新生成、对比不同版本的效果。'),
]

for idx, text in steps_42:
    set_text(doc.paragraphs[idx], text)

print('OK §4.2 naturalized')

# ====================================================================
# §5.1 技术架构 — tone down from dev docs
# ====================================================================
set_text(doc.paragraphs[91],
    '后端采用 FastAPI 框架，按分层架构组织：'
    '接口层处理 HTTP 请求路由和参数校验；'
    '应用层编排业务用例——提交题目生成动画、导出视频、管理账户等；'
    '领域层封装核心模型（教学蓝图、教学脚本、执行映射等数据结构）'
    '和领域服务（蓝图生成提示词构建、脚本装配映射、学科路由分类、'
    '数学几何校验等纯函数计算）；'
    '基础设施层实现具体的大模型 HTTP 调用、'
    'SQLite 数据持久化、微信支付接入等。'
    '学科技能包（如代数、微积分、立体几何、力学等）以声明式注册，'
    '新学科可通过配置文件零代码扩展。'
    '\n\n'
    '前端采用 React 19 和 TypeScript，'
    '以 Remotion 渲染引擎为核心实现帧精确动画播放。'
    '代码按功能模块分层组织：'
    '共享组件库、实体类型定义、'
    '播放引擎（渲染器注册表 + 播放器 + 参数面板 + 语音合成 + 导出）、'
    '以及页面层（输入页、工作室页、历史页）。'
    '交互播放和视频导出共用同一套渲染器注册表，确保画面一致。'
    '\n\n'
    '智能体模式（可选启用）：独立的 Node.js 侧边车服务'
    '提供 14 个底层绘图工具和 11 个跨学科教学模板。'
    '大模型通过逐步调用这些工具精确构建动画——'
    '规划大纲、放置数组标记、添加函数曲线、写入公式——'
    '每一步都经过几何自检，最终由独立评审模型复核。'
    '\n\n'
    '教学脚本是平台唯一的数据契约（JSON 格式）。'
    '只描述教学内容而不绑定渲染技术，'
    '意味着同一份脚本可以驱动不同的输出端——'
    '当前稳定支持交互播放和视频导出，未来可扩展至更多形态。')

print('OK §5.1 toned down')

# ====================================================================
# §5.2 核心算法流程 — clean
# ====================================================================
set_text(doc.paragraphs[93],
    'MetaView 提供两种生成模式，用户可按需选择：'
    '\n\n'
    '标准模式（默认，速度快，适合日常教学）：'
    '\n'
    '题目输入 → 学科路由自动分类（混合策略：小模型快速判断 + 规则兜底）'
    '→ 大模型按 JSON 格式规范输出教学蓝图和执行映射'
    '→ 自动校验（结构完整性、字段一致性、讲解词质量）'
    '→ 不通过则自动修复或让模型重试（最多可配置 N 次）'
    '→ 校验通过后装配为教学脚本'
    '→ 前端渲染引擎逐帧播放，或导出为视频文件。'
    '\n\n'
    '智能体模式（速度较慢，精细度高，适合高质量教学视频制作）：'
    '\n'
    '题目输入 → 大模型通过绘图工具链逐步构建——'
    '先规划整体大纲，再逐步骤添加曲线、数组标记、公式等元素，'
    '每一步完成后进行几何自检（方向判断、经过点判断、单调性判断，'
    '由后端的数学计算引擎给出确定真值），'
    '全部步骤完成后经自我检查和独立评审模型复核，'
    '再通过渲染器兼容性验证，最终输出教学脚本。'
    '\n\n'
    '两条路径输出完全相同的教学脚本格式，走完全相同的渲染出口，'
    '差别仅在于生成过程的控制粒度。')

print('OK §5.2 cleaned')

# ====================================================================
# §5.3 创新点 — clean
# ====================================================================
set_text(doc.paragraphs[95],
    '创新一：教学蓝图——语言无关的结构化教学契约。'
    '模型只输出符合严格规范的 JSON 数据（标题、步骤、视觉类型、'
    '标记点、讲解词模板、函数图参数等），'
    '画面渲染完全由系统控制而非模型生成。'
    '这从架构层面阻断了长文本截断和幻觉扩散——'
    '不合规的输出在管线内被拦截修复，不会到达用户。'
    '\n\n'
    '创新二：双模式共享唯一出口。'
    '标准模式走快速生成路径，智能体模式通过绘图工具链逐步精细构建——'
    '两种模式输出同一格式的教学脚本，走同一套渲染器，'
    '用户可根据场景灵活选择，不被锁定在单一模型或供应商上。'
    '\n\n'
    '创新三：自动校验-修复-重试质量闭环。'
    '多维自动检查（结构 / 标记 / 视觉类型 / 讲解词质量）→ '
    '自动修复缺失字段和布局问题 → '
    '修复失败则反馈大模型重新生成。主动修复替代报错。'
    '\n\n'
    '创新四：大小模型分层降低调用成本。'
    '低成本小模型（如 DeepSeek、gpt-4o-mini）做学科路由和意图分类，'
    '高质量大模型做教学蓝图生成。'
    '学科技能包声明式注册，新学科零代码适配，'
    'API 调用成本相比全量大模型方案降低 30% 以上。'
    '\n\n'
    '创新五：代码-动画-时间轴-数据状态四向联动。'
    '执行映射将代码行号、动画步骤、时间轴位置、变量状态四者精确对齐，'
    '用户点击任意维度即刻跳转到对应位置——'
    '从"看动画"升级为"探索过程"。'
    '\n\n'
    '创新六：绘图工具链——让大模型一步步精确构建。'
    '14 个底层原子工具 + 11 个跨学科教学模板 + 数学引擎几何自检'
    '（方向性 / 经过点 / 单调性判断），'
    '从工具能力面消除大模型的无意义铺陈，每一帧都有据可查。')

print('OK §5.3 cleaned')

# ====================================================================
# TABLES: sweep prohibited terms
# ====================================================================
# Map of (table_finder_text, row_col_pairs) → new text
# We identify tables by first-row content, then fix specific cells

table_fixes = {}

# Walk all tables and apply fixes based on cell text matching
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if not txt:
                continue

            # Fix patterns
            new_txt = None

            # 三维模型 → 交互动画 (unless describing competitors)
            if '三维模型' in txt and 'VR/AR' not in txt and '虚拟实验、VR' not in txt:
                new_txt = txt.replace('三维模型', '交互动画')

            # 虚拟实验 → remove (unless competitor context)
            if '虚拟实验' in txt and 'VR/AR' not in txt and '虚拟实验、VR' not in txt:
                new_txt = (new_txt or txt).replace('、虚拟实验', '').replace('，虚拟实验', '').replace('虚拟实验、', '').replace('虚拟实验', '参数可调模拟')

            # 知识图谱 → 学科知识体系
            if '知识图谱' in txt:
                new_txt = (new_txt or txt).replace('知识图谱', '学科知识体系')

            # RAG / 检索增强生成
            if 'RAG' in txt or '检索增强生成' in txt:
                new_txt = (new_txt or txt).replace('检索增强生成（RAG）', '提示词工程').replace('RAG', '提示词工程')

            # 来源追踪 → remove or soften
            if '来源追踪' in txt:
                new_txt = (new_txt or txt).replace('、来源追踪', '').replace('来源追踪和', '').replace('来源追踪', '校准校验')

            # 专家审校 → 教师审校（可选环节）
            if '专家审校' in txt:
                new_txt = (new_txt or txt).replace('专家审校', '教师审校')

            # 版权标注 → remove
            if '版权标注' in txt:
                new_txt = (new_txt or txt).replace('、版权标注', '').replace('版权标注、', '').replace('版权标注', '授权标记')

            # 敏感词过滤 / 内容审核 → 基本输入校验
            if '敏感词过滤' in txt or '敏感信息过滤' in txt:
                new_txt = (new_txt or txt).replace('、敏感词和风险内容过滤', '').replace('敏感信息过滤、', '').replace('敏感信息过滤', '基本校验')

            # 内容审核 → soften
            if '内容审核' in txt and '竞品' not in txt:
                new_txt = (new_txt or txt).replace('、内容审核', '').replace('内容审核、', '').replace('内容审核', '内容校验')

            # 学生学习空间 (not in future context)
            if '学生学习空间' in txt and '规划' not in txt:
                new_txt = (new_txt or txt).replace('、学生学习空间', '（规划中）').replace('学生学习空间、', '')

            # 创作者工作台 (not built)
            if '创作者工作台' in txt and '规划' not in txt:
                new_txt = (new_txt or txt).replace('、科普创作者工作台', '').replace('科普创作者工作台、', '').replace('科普创作者工作台', '内容导出工具（规划中）')

            # 学校管理后台 (not built)
            if '学校管理后台' in txt and '规划' not in txt:
                new_txt = (new_txt or txt).replace('、校级管理与内容审核后台', '（规划中）').replace('校级管理与内容审核后台、', '')

            # 校本资源库 (not built, mark as planned)
            if '校本资源库' in txt and '规划' not in txt:
                new_txt = (new_txt or txt).replace('、校本资源库', '（规划中）').replace('校本资源库、', '').replace('校本资源库', '教学资源沉淀（规划中）')

            # PPT/图片/动图/HTML5 output types → fix
            if 'PPT/图片/动图/HTML5' in txt or 'PPT/图片' in txt:
                new_txt = (new_txt or txt).replace('PPT/图片/动图/HTML5', '交互播放器 / 视频文件').replace('PPT/图片导出', '视频文件导出').replace('PPT/图片', '交互播放 / 视频文件')

            # 六部分组成 → 三层架构
            if '六部分组成' in txt:
                new_txt = (new_txt or txt).replace('六部分组成', '三个服务组成')

            # 教师端、学生端、创作者端 → fix
            if '教师端、学生端、创作者端' in txt:
                new_txt = (new_txt or txt).replace('教师端、学生端、创作者端、学校管理端、数据看板、导出接口',
                                                  'Web 前端、API 后端、数据看板、导出接口')

            # 教师端和学生端 in vision context
            if '教师端和' in txt and 'AI 可视化引擎' in txt:
                new_txt = (new_txt or txt).replace('教师端和 AI 可视化引擎', 'Web 前端和 API 后端')

            if new_txt and new_txt != txt:
                set_cell(table, ri, ci, new_txt)

# Special: Table 11 (输出形态) — complete rewrite
# Find it by the "课堂课件" cell
for table in doc.tables:
    for ri, row in enumerate(table.rows):
        if row.cells[0].text.strip() == '课堂课件':
            set_cell(table, 1, 0, '交互播放')
            set_cell(table, 1, 1, '帧驱动播放器 / 步骤跳转 / 参数面板 / 字幕叠加')
            set_cell(table, 1, 2, '适合课堂演示和自主探索。')
            set_cell(table, 2, 0, '视频导出')
            set_cell(table, 2, 1, 'MP4 / WebM / GIF / 可选语音合成配音')
            set_cell(table, 2, 2, '适合微课制作、预习复习、内容分发。')
            set_cell(table, 3, 0, '四向联动回放')
            set_cell(table, 3, 1, '代码行高亮 + 动画步骤 + 时间轴 + 数据状态联动')
            set_cell(table, 3, 2, '适合编程教学和算法讲解。')
            set_cell(table, 4, 0, '数学函数图')
            set_cell(table, 4, 1, '坐标系曲线绘制 / 导数切线 / 积分阴影 / 公式渲染')
            set_cell(table, 4, 2, '适合数学和理工科教学。')
            set_cell(table, 5, 0, '教学脚本数据')
            set_cell(table, 5, 1, '结构化 JSON / 运行历史 / 原始题目回溯')
            set_cell(table, 5, 2, '适合资源沉淀、复盘分析和批量处理。')
            print('OK Table 11 rewritten')
            break

# Also fix Table 32 (风险控制) - RAG and related
for table in doc.tables:
    for ri, row in enumerate(table.rows):
        if row.cells[0].text.strip() == '技术风险':
            if 'RAG' in row.cells[2].text:
                set_cell(table, ri, 2,
                    '蓝图校验器多维自动检查 + 修复服务自动补全 + 大模型重试重生成。'
                    '校验规则可配置，不通过不交付。')
                print('OK Table 32 (技术风险)')
        if row.cells[0].text.strip() == '内容风险':
            if '来源追踪' in row.cells[2].text or '专家审校' in row.cells[2].text:
                set_cell(table, ri, 2,
                    '自动校验 + 自动修复 + 大模型重试闭环。'
                    '关键教学案例可经教师审校确认（可选环节）。'
                    '数学表达式经白名单字符过滤防注入。')
                print('OK Table 32 (内容风险)')

# ====================================================================
# SAVE
# ====================================================================
doc.save(DST)
print(f'\nDone → {DST}')
