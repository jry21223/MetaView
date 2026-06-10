"""
Fix the copy document: formatting + content from V2 + table cleanup.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy

SRC = '/Users/jerry/Desktop/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2（副本）.docx'
V2 = '/Users/jerry/Downloads/MetaView_商业计划书_中国国际大学生创新大赛参赛版_v2.docx'
DST = SRC  # overwrite in place

doc = Document(SRC)
doc_v2 = Document(V2)


def set_text(para, text):
    runs = para.runs
    if not runs:
        r = para.add_run(text)
        r.font.name = 'Noto Sans CJK SC'
        r.font.size = Pt(11)
        return
    for r in runs:
        r.text = ''
    runs[0].text = text


def set_cell(table, row, col, text):
    cell = table.rows[row].cells[col]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs:
        rr = cell.paragraphs[0].runs
        if rr: rr[0].text = text
        else: cell.paragraphs[0].text = text
    else:
        cell.text = text


# ============================================================
# PHASE 1: Copy cleaned content from V2 into this document
# ============================================================

# Find the cleaned paragraphs in V2 by their content signatures
v2_texts = {}  # signature → full text
for p in doc_v2.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    # Store signature based on first 40 chars
    sig = txt[:40].replace('\n', ' ')
    if sig not in v2_texts:
        v2_texts[sig] = txt

# --- §4.2 steps: find V2 versions ---
# V2 step signatures
steps_v2_sigs = [
    '1. 输入题目。用户在输入页面以自然语言、代码片段或数学公式描述题目。如果涉及编程教学，可以粘贴源代码，系统会自动关联代码行和动画步骤。',
    '2. 自动识别学科。系统根据题目内容自动判断所属学科领域（当前覆盖算法、数学、代码、物理、化学、生物、地理七个方向），并匹配合适的教学策略和步骤规划。',
    '3. 生成教学蓝图。大语言模型根据题目和学科，输出一份结构化的教学方案：',
    '4. 自动校验与装配。系统对生成的教学方案进行多维检查——结构是否完整、标记是否一致、讲解是否达标。',
    '5. 交互播放与导出。前端渲染引擎按每秒 30 帧驱动教学动画。',
    '6. 历史沉淀与复盘。每次生成自动保存原始题目和完整脚本。',
]

# Find which paragraphs in COPY should be replaced (P76-P81)
# and replace with V2 versions
for idx, v2_text in zip(range(76, 82), steps_v2_sigs):
    if idx < len(doc.paragraphs):
        # Find matching text in V2
        for p in doc_v2.paragraphs:
            if p.text.strip().startswith(v2_text[:30]):
                set_text(doc.paragraphs[idx], p.text.strip())
                print(f'OK §4.2 P{idx}: V2 content applied')
                break
        else:
            # Fallback: use the sig text directly
            set_text(doc.paragraphs[idx], v2_text)
            print(f'OK §4.2 P{idx}: fallback text applied')

# --- §5.2: algorithm flow ---
v2_52_sig = 'MetaView 提供两种生成模式，用户可按需选择：'
for p in doc_v2.paragraphs:
    if p.text.strip().startswith(v2_52_sig):
        # Find the corresponding paragraph in COPY (P91)
        for i, cp in enumerate(doc.paragraphs):
            if cp.text.strip().startswith('MetaView v2 支持两条生成路径'):
                set_text(doc.paragraphs[i], p.text.strip())
                print(f'OK §5.2 P{i}: V2 content applied')
                break
        break

# --- §5.3: innovation points ---
v2_53_sig = '创新一：教学蓝图——语言无关的结构化教学契约。'
for p in doc_v2.paragraphs:
    if p.text.strip().startswith(v2_53_sig):
        for i, cp in enumerate(doc.paragraphs):
            if cp.text.strip().startswith('创新 1'):
                set_text(doc.paragraphs[i], p.text.strip())
                print(f'OK §5.3 P{i}: V2 content applied')
                break
        break

# --- §5.4: first bullet (来源追踪) ---
for i, cp in enumerate(doc.paragraphs):
    if cp.text.strip().startswith('来源追踪：每个知识卡片'):
        set_text(doc.paragraphs[i],
            '自动校验闭环：每次生成的教学脚本都经过蓝图校验器多维检查'
            '（结构完整性、字段一致性、视觉类型合法性、讲解词质量），'
            '不通过的输出在管线内被拦截修复，不会到达用户。'
            '修复失败则触发大模型重新生成，形成闭环保障。')
        print(f'OK §5.4 P{i}: 来源追踪 → 自动校验闭环')
        break

# --- P89: remove "PPT" from output list ---
for i, p in enumerate(doc.paragraphs):
    if '驱动交互播放器、视频导出、PPT、HTML课件' in p.text:
        for run in p.runs:
            run.text = run.text.replace('PPT、HTML课件', 'HTML 课件')
        print(f'OK P{i}: removed PPT from output list')
        break

# ============================================================
# PHASE 2: Formatting — unify fonts, sizes, margins
# ============================================================
HEADING_FONT = 'Noto Serif CJK SC'
BODY_FONT = 'Noto Sans CJK SC'

for para in doc.paragraphs:
    style = para.style.name if para.style else ''

    if style == 'Heading 1':
        for run in para.runs:
            run.font.name = HEADING_FONT
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    elif style == 'Heading 2':
        for run in para.runs:
            run.font.name = HEADING_FONT
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    elif style in ('Normal', 'List Bullet'):
        for run in para.runs:
            if run.font.name and 'Noto' not in run.font.name:
                run.font.name = BODY_FONT
            if run.font.size and run.font.size > Pt(12):
                pass  # keep intentional large sizes
            elif not run.font.size:
                run.font.size = Pt(11)

# Tables: unify font
for table in doc.tables:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if r_idx == 0:
                        run.font.name = HEADING_FONT
                        run.font.size = Pt(10)
                        run.font.bold = True
                    else:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(9.5)

# Remove excessive empty paragraphs
# Strategy: if 3+ consecutive empty Normal paragraphs, keep only 1
i = 0
while i < len(doc.paragraphs) - 2:
    p = doc.paragraphs[i]
    if p.style.name == 'Normal' and p.text.strip() == '':
        # Check if next 2 are also empty
        if (i+1 < len(doc.paragraphs) and
            doc.paragraphs[i+1].style.name == 'Normal' and
            doc.paragraphs[i+1].text.strip() == '' and
            i+2 < len(doc.paragraphs) and
            doc.paragraphs[i+2].style.name == 'Normal' and
            doc.paragraphs[i+2].text.strip() == ''):
            # Remove paragraph at i+1 (merge 3 into 2 by removing the middle one)
            p_elem = doc.paragraphs[i+1]._element
            p_elem.getparent().remove(p_elem)
            continue  # re-check at same position
    i += 1

print('OK formatting unified')

# ============================================================
# PHASE 3: Table cleanup
# ============================================================
fixes = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if not txt:
                continue
            new_txt = txt

            # 三维模型 → 交互动画 (skip competitor analysis)
            if '三维模型' in txt and 'VR/AR' not in txt:
                new_txt = new_txt.replace('三维模型', '交互动画')
            if '三维结构' in txt:
                new_txt = new_txt.replace('三维结构', '空间结构')

            # 虚拟实验 → remove (skip competitor)
            if '虚拟实验' in txt and 'VR/AR' not in txt:
                new_txt = new_txt.replace('、虚拟实验', '').replace('虚拟实验、', '').replace('虚拟实验', '参数可调模拟')

            # 知识图谱 → 学科知识体系
            if '知识图谱' in txt:
                new_txt = new_txt.replace('知识图谱', '学科知识体系')

            # RAG
            if 'RAG' in txt:
                new_txt = new_txt.replace('、RAG', '').replace('RAG、', '').replace('RAG', '提示词工程')

            # 来源追踪
            if '来源追踪' in txt:
                new_txt = new_txt.replace('、来源追踪', '').replace('来源追踪、', '').replace('来源追踪', '')

            # 版权标注
            if '版权标注' in txt:
                new_txt = new_txt.replace('、版权标注', '').replace('版权标注、', '').replace('版权标注', '授权标记')

            # 敏感词
            if '敏感词' in txt:
                new_txt = new_txt.replace('、敏感词和风险内容过滤', '').replace('敏感词和风险内容过滤、', '')

            # 专家审校 → 教师审校
            if '专家审校' in txt:
                new_txt = new_txt.replace('专家审校', '教师审校')

            # 学生学习空间
            if '学生学习空间' in txt:
                new_txt = new_txt.replace('、学生学习空间', '').replace('学生学习空间', '交互数学画板')

            # 创作者工作台
            if '创作者工作台' in txt:
                new_txt = new_txt.replace('、科普创作者工作台', '（规划中）').replace('科普创作者工作台', '内容导出工具（规划中）')

            # PPT/图片
            if 'PPT/图片' in txt:
                if 'PPT/图片/动图/HTML5' in txt:
                    new_txt = new_txt.replace('PPT/图片/动图/HTML5', '交互播放器 / 视频文件')
                else:
                    new_txt = new_txt.replace('PPT/图片', '交互播放 / 视频文件')
            if 'MP4/PPT/图片' in new_txt:
                new_txt = new_txt.replace('MP4/PPT/图片', 'MP4/视频文件/图片')

            # 内容审核 → soften (but NOT in competitor tables)
            if '内容审核' in txt and '竞品' not in ' '.join(table.rows[0].cells[0].text):
                new_txt = new_txt.replace('、内容审核', '').replace('内容审核、', '').replace('内容审核', '内容校验')

            if new_txt != txt:
                set_cell(table, ri, ci, new_txt)
                fixes += 1

print(f'OK tables: {fixes} cell fixes applied')

# ============================================================
# SAVE
# ============================================================
doc.save(DST)
print(f'\nDone → {DST}')
